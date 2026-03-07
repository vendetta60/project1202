from app.repositories.report import ReportRepository
from app.schemas.report import ReportResponse, ReportItem, ReportParams
from app.models.user import User
import pandas as pd
import io
from datetime import date, datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

# Try to register a font that supports Azerbaijani characters
try:
    font_path = "C:\\Windows\\Fonts\\arial.ttf"
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont('Arial', font_path))
        DEFAULT_FONT = 'Arial'
    else:
        DEFAULT_FONT = 'Helvetica'
except Exception:
    DEFAULT_FONT = 'Helvetica'

class ReportService:
    def __init__(self, reports: ReportRepository):
        self.reports = reports

    def get_appeal_report(self, params: ReportParams, user: User) -> ReportResponse:
        user_section_id = None
        if not user.is_admin:
            user_section_id = user.section_id

        results = self.reports.get_appeal_stats(
            group_by=params.group_by,
            start_date=params.start_date,
            end_date=params.end_date,
            user_section_id=user_section_id
        )

        items = []
        total_count = 0
        for row in results:
            count = row.count
            total_count += count
            items.append(ReportItem(
                id=row.id,
                name=row.name or "Məlum deyil",
                count=count
            ))

        return ReportResponse(
            items=items,
            total=total_count,
            group_by=params.group_by,
            start_date=params.start_date,
            end_date=params.end_date
        )

    def _prepare_forma_4_rows(self, results):
        rows = []
        for ap in results:
            # Column 5: haradan gəlib
            haradan_parts = []
            if ap.department:
                haradan_parts.append(ap.department.department)
            if ap.official_rel:
                haradan_parts.append(ap.official_rel.official)
            haradan = "\n".join(haradan_parts)
            
            # Column 11: Kim baxmışdır və dərkənar
            derkenar_parts = []
            if ap.control_rel:
                derkenar_parts.append(ap.control_rel.chief or "")
            if ap.instruction_rel:
                derkenar_parts.append(ap.instruction_rel.instructions or "")
            derkenar = "\n".join(derkenar_parts)
            
            # Executors data
            execs = ap.executors or []
            directions = "\n".join(set(e.direction_name for e in execs if e.direction_name))
            executor_names = "\n".join(set(e.executor_name for e in execs if e.executor_name))
            
            # Col 14: out_num and out_date
            icra_senedi_parts = []
            for e in execs:
                if e.out_num:
                    s = str(e.out_num)
                    if e.out_date:
                        s += f" {e.out_date.strftime('%d.%m.%Y')}"
                    icra_senedi_parts.append(s)
            icra_senedi = "\n".join(icra_senedi_parts)

            # Col 18: r_num and r_date
            gonderilme_parts = []
            for e in execs:
                if e.r_num:
                    s = str(e.r_num)
                    if e.r_date:
                        s += f" {e.r_date.strftime('%d.%m.%Y')}"
                    gonderilme_parts.append(s)
            gonderilme = "\n".join(gonderilme_parts)

            # Col 16-17: Tikildiyi iş №-si and İşdəki vərəq №-si
            tikildiyi_is_parts = []
            isdeki_vereq_parts = []
            for e in execs:
                if getattr(e, "attach_num", None):
                    tikildiyi_is_parts.append(str(e.attach_num))
                if getattr(e, "attach_paper_num", None):
                    isdeki_vereq_parts.append(str(e.attach_paper_num))
            tikildiyi_is = "\n".join(tikildiyi_is_parts)
            isdeki_vereq = "\n".join(isdeki_vereq_parts)

            # Column 10: müraciətin növü + təkrar olub-olmaması
            content_type_parts = []
            if ap.content_type_rel and ap.content_type_rel.content_type:
                content_type_parts.append(ap.content_type_rel.content_type)
            if getattr(ap, "repetition", None):
                content_type_parts.append("Təkrar")
            content_type_value = "\n".join(content_type_parts)

            row = {
                "1": ap.reg_num or "",
                "2": ap.reg_date.strftime("%d.%m.%Y") if ap.reg_date else "",
                "3": ap.in_ap_num or "",
                "4": ap.in_ap_date.strftime("%d.%m.%Y") if ap.in_ap_date else "",
                "5": haradan,
                "6": ap.content or "",
                # Müraciətin indeksi sahəsində tam ad yox, yalnız indeks nömrəsi göstərilsin
                "7": str(ap.ap_index_rel.ap_index_id) if ap.ap_index_rel and ap.ap_index_rel.ap_index_id is not None else "",
                "8": ap.paper_count or "",
                # Hesabat indeksi sahəsində də yalnız indeks nömrəsi göstərilsin
                "9": str(ap.account_index_id) if getattr(ap, "account_index_id", None) is not None else "",
                "10": content_type_value,
                "11": derkenar,
                "12": directions,
                "13": executor_names,
                "14": icra_senedi,
                "15": ap.status_rel.status if ap.status_rel else "",
                "16": tikildiyi_is,
                "17": isdeki_vereq,
                "18": gonderilme
            }
            rows.append(row)
        return rows

    def generate_forma_4_excel(self, start_date: date | None, end_date: date | None, user: User) -> io.BytesIO:
        user_section_id = None if user.is_admin else user.section_id
        results = self.reports.get_forma_4_data(start_date, end_date, user_section_id)

        headers = [
            "Qeydəalınma №-si", "Qeydəalınma tarixi", "Daxil olan müraciətin №-si", "Daxil olan müraciətin tarixi",
            "Müraciət haradan (kimdən) gəlib", "Müraciətin qısa məzmunu", "Müraciətin indeksi", "Vərəq sayı",
            "Hesabat indeksi", "Müraciətin növü", "Kim baxmışdır və dərkənar",
            "Müraciət hansı struktur bölməyə icraya verilib", "İcraçının adı və soyadı",
            "Müraciət hansı sənədlə icra edilib", "Müraciətin baxılması vəziyyəti",
            "Tikildiyi iş №-si", "İşdəki vərəq №-si", "Sənədin göndərilməsi barədə qeyd"
        ]

        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Side, Font

        wb = Workbook()
        ws = wb.active
        ws.title = "Forma 4"

        # Header rows
        ws.append(headers)
        ws.append(list(range(1, 19)))

        # Data rows start from row 3
        if results:
            prepared_rows = self._prepare_forma_4_rows(results)
            for r in prepared_rows:
                ws.append([r.get(str(i), "") for i in range(1, 19)])
        else:
            ws.append([""] * 18)

        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        col_widths = [15, 12, 15, 12, 25, 30, 10, 8, 10, 15, 25, 20, 20, 20, 15, 12, 12, 15]
        for i, width in enumerate(col_widths, start=1):
            ws.column_dimensions[chr(64 + i)].width = width

        ws.row_dimensions[1].height = 120
        ws.row_dimensions[2].height = 22
        ws.freeze_panes = "A3"

        header_font = Font(bold=True, size=10)
        normal_font = Font(size=10)

        max_row = ws.max_row
        max_col = 18

        for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
            for cell in row:
                cell.border = thin_border
                if cell.row == 1:
                    cell.font = header_font
                    cell.alignment = Alignment(
                        textRotation=90,
                        wrap_text=True,
                        vertical="center",
                        horizontal="center",
                    )
                elif cell.row == 2:
                    cell.font = header_font
                    cell.alignment = Alignment(
                        wrap_text=True,
                        vertical="center",
                        horizontal="center",
                    )
                else:
                    cell.font = normal_font
                    cell.alignment = Alignment(
                        wrap_text=True,
                        vertical="top",
                        horizontal="center",
                    )

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    def generate_forma_4_word(self, start_date: date | None, end_date: date | None, user: User) -> io.BytesIO:
        user_section_id = None if user.is_admin else user.section_id
        results = self.reports.get_forma_4_data(start_date, end_date, user_section_id)
        rows = self._prepare_forma_4_rows(results)

        doc = Document()

        # Page setup: A4 landscape (album forması)
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width / height to apply orientation
        section.page_width, section.page_height = section.page_height, section.page_width
        # Tighter margins so cədvəl A4-ə yaxşı sığsın
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)

        # Base font for whole document
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(9)

        title_para = doc.add_paragraph("Daxil olan vətəndaş müraciətlərinin qeydiyyatı JURNALI")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.size = Pt(12)

        # Column headers – eyni Excel forması ilə
        headers = [
            "Qeydəalınma №-si", "Qeydəalınma tarixi", "Daxil olan müraciətin №-si",
            "Daxil olan müraciətin tarixi", "Müraciət haradan (kimdən) gəlib",
            "Müraciətin qısa məzmunu", "Müraciətin indeksi", "Vərəq sayı",
            "Hesabat indeksi", "Müraciətin növü", "Kim baxmışdır və dərkənar",
            "Müraciət hansı struktur bölməyə icraya verilib", "İcraçının adı və soyadı",
            "Müraciət hansı sənədlə icra edilib", "Müraciətin baxılması vəziyyəti",
            "Tikildiyi iş №-si", "İşdəki vərəq №-si", "Sənədin göndərilməsi barədə qeyd",
        ]

        table = doc.add_table(rows=2, cols=18)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 1-ci sətr: başlıq mətnləri (qalın, mərkəzdə)
        header_row = table.rows[0]
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)

        # 2-ci sətr: sütun nömrələri 1–18
        number_row = table.rows[1]
        for i in range(18):
            number_row.cells[i].text = str(i + 1)
            for p in number_row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)

        if not results:
            data_row = table.add_row()
            data_row.cells[0].text = "Məlumat tapılmadı"
        else:
            for r in rows:
                row_cells = table.add_row().cells
                for i in range(18):
                    row_cells[i].text = str(r.get(str(i + 1), "")) or ""

        # Sütun enləri – Excel col_widths-dən təxmini çevrilmə
        col_widths = [15, 12, 15, 12, 25, 30, 10, 8, 10, 15, 25, 20, 20, 20, 15, 12, 12, 15]
        for i, width in enumerate(col_widths):
            # 1 Excel vahidini təxminən 0.12 inch kimi götürək
            table.columns[i].width = Inches(width * 0.12)

        # Hüceyrə align və vertical align
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def generate_forma_4_pdf(self, start_date: date | None, end_date: date | None, user: User) -> io.BytesIO:
        user_section_id = None if user.is_admin else user.section_id
        results = self.reports.get_forma_4_data(start_date, end_date, user_section_id)
        rows = self._prepare_forma_4_rows(results)
        
        output = io.BytesIO()
        doc = SimpleDocTemplate(output, pagesize=landscape(A4))
        elements = []
        
        styles = getSampleStyleSheet()
        # Başlıq – Forma 4-ün üst hissəsi
        title_style = styles['Title']
        title_style.fontName = DEFAULT_FONT
        elements.append(Paragraph("Daxil olan vətəndaş müraciətlərinin qeydiyyatı JURNALI", title_style))
        elements.append(Spacer(1, 12))
        
        # Sütun başlıqları – Excel və Word-dəki kimi
        headers = [
            "Qeydəalınma №-si", "Qeydəalınma tarixi", "Daxil olan müraciətin №-si",
            "Daxil olan müraciətin tarixi", "Müraciət haradan (kimdən) gəlib",
            "Müraciətin qısa məzmunu", "Müraciətin indeksi", "Vərəq sayı",
            "Hesabat indeksi", "Müraciətin növü", "Kim baxmışdır və dərkənar",
            "Müraciət hansı struktur bölməyə icraya verilib", "İcraçının adı və soyadı",
            "Müraciət hansı sənədlə icra edilib", "Müraciətin baxılması vəziyyəti",
            "Tikildiyi iş №-si", "İşdəki vərəq №-si", "Sənədin göndərilməsi barədə qeyd",
        ]

        # 1-ci sətr: mətn başlıqları, 2-ci sətr: sütun nömrələri
        data = [headers, ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12", "13", "14", "15", "16", "17", "18"]]
        if not results:
            data.append(["Məlumat tapılmadı"] + [""] * 17)
        else:
            for r in rows:
                data.append([str(r[str(i+1)]) for i in range(1, 19)])
            
        t = Table(data, repeatRows=2)
        t.setStyle(TableStyle([
            # Başlıq sətrləri
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('BACKGROUND', (0, 1), (-1, 1), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTNAME', (0, 0), (-1, 1), f'{DEFAULT_FONT}-Bold'),
            ('FONTNAME', (0, 2), (-1, -1), DEFAULT_FONT),
            ('FONTSIZE', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, 1), 6),
            ('BACKGROUND', (0, 2), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
        ]))
        
        elements.append(t)
        doc.build(elements)
        output.seek(0)
        return output

    def _get_group_label(self, group_by: str) -> str:
        labels = {
            "department": "İdarə",
            "region": "Region",
            "status": "Müraciətin baxılması vəziyyəti",
            "index": "Müraciətin indeksi",
            "account_index": "Hesabat indeksi",
            "exec_direction": "İcra struktur bölmə",
            "content_type": "Müraciətin növü",
            "insection": "Hərbi hissə (Bölmə)"
        }
        return labels.get(group_by, "Kateqoriya")

    def generate_appeal_stats_excel(self, params: ReportParams, user: User) -> io.BytesIO:
        report = self.get_appeal_report(params, user)
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Side, Font
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Hesabat"
        
        group_label = self._get_group_label(params.group_by)
        ws.append(["MÜRACİƏTLƏRİN STATİSTİK HESABATI"])
        ws.append([f"Qruplaşdırma: {group_label} üzrə"])
        ws.append([f"Dövr: {params.start_date or 'Əvvəldən'} - {params.end_date or 'Bugünədək'}"])
        ws.append([])
        
        headers = [group_label, "Sayı", "Nisbət (%)"]
        ws.append(headers)
        
        for item in report.items:
            percentage = (item.count / report.total * 100) if report.total > 0 else 0
            ws.append([item.name, item.count, f"{percentage:.1f}%"])
            
        ws.append([])
        ws.append(["CƏMİ", report.total, "100.0%"])
        
        # Styling
        header_font = Font(bold=True)
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        
        for cell in ws[5]:
            cell.font = header_font
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
            
        for row in ws.iter_rows(min_row=6, max_row=ws.max_row):
            for cell in row:
                cell.border = thin_border
                
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output

    def generate_appeal_stats_word(self, params: ReportParams, user: User) -> io.BytesIO:
        report = self.get_appeal_report(params, user)
        group_label = self._get_group_label(params.group_by)
        period = f"{params.start_date or 'Əvvəldən'} — {params.end_date or 'Bugünədək'}"

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = doc.add_paragraph("MÜRACİƏTLƏRİN STATİSTİK HESABATI")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.bold = True
        title.runs[0].font.size = Pt(14)

        doc.add_paragraph(f"Qruplaşdırma: {group_label} üzrə")
        doc.add_paragraph(f"Hesabat dövrü: {period}")
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = group_label
        hdr[1].text = "Sayı"
        hdr[2].text = "Nisbət (%)"
        for c in hdr:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True

        for item in report.items:
            row = table.add_row().cells
            pct = (item.count / report.total * 100) if report.total > 0 else 0
            row[0].text = item.name
            row[1].text = str(item.count)
            row[2].text = f"{pct:.1f}%"
            for c in row:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        total_row = table.add_row().cells
        total_row[0].text = "CƏMİ"
        total_row[1].text = str(report.total)
        total_row[2].text = "100.0%"
        for c in total_row:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in c.paragraphs[0].runs:
                run.font.bold = True

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def generate_appeal_stats_pdf(self, params: ReportParams, user: User) -> io.BytesIO:
        report = self.get_appeal_report(params, user)
        output = io.BytesIO()
        doc = SimpleDocTemplate(output, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        group_label = self._get_group_label(params.group_by)
        
        elements.append(Paragraph("MÜRACİƏTLƏRİN STATİSTİK HESABATI", styles['Title']))
        elements.append(Paragraph(f"Qruplaşdırma: {group_label} üzrə", styles['Normal']))
        elements.append(Paragraph(f"Dövr: {params.start_date or 'Əvvəldən'} - {params.end_date or 'Bugünədək'}", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        data = [[group_label, "Sayı", "Nisbət (%)"]]
        for item in report.items:
            percentage = (item.count / report.total * 100) if report.total > 0 else 0
            data.append([item.name, str(item.count), f"{percentage:.1f}%"])
            
        data.append(["CƏMİ", str(report.total), "100.0%"])
        
        t = Table(data, hAlign='LEFT')
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), f'{DEFAULT_FONT}-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), DEFAULT_FONT),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
            ('FONTNAME', (0, -1), (-1, -1), f'{DEFAULT_FONT}-Bold'),
        ]))
        
        elements.append(t)
        doc.build(elements)
        output.seek(0)
        return output
